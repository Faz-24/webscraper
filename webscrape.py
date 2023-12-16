import os 
import requests
from bs4 import BeautifulSoup
from docx import Document
import time

# URL of the webpage to scrape
urls = ["https://www.novelcool.com/chapter/Damn-Reincarnation-Chapter-325/11426653/"]
counter = 324 
while len(urls) > 0:
# Send an HTTP GET request to the URL
    counter += 1
    
    response = requests.get(urls[0])
    time.sleep(5)
    if response.status_code == 200:
        # Parse the HTML content of the page
        soup = BeautifulSoup(response.text, 'html.parser')

        # Find all the <relevant div tags> tags on the page that contain the p tags
        div_tags = soup.find('div', {'class': 'overflow-hidden'})

        if div_tags:
            p_tag = div_tags.get_text(separator="\n \n")
            #p_tags = div_tags.find('p')                  
            #p_tag = p_tags.get_text(separator="\n")

        # Specify the destination folder path
        destination_folder = "/Users/f.shahriyar/Documents"

        # Create a new Word document
        doc = Document()
        doc.add_paragraph(p_tag)
        doc.add_paragraph("fahim")
        # Specify the full path to the output file in the destination folder
        output_file = os.path.join(destination_folder, f"output-{counter}.docx")

        # Save the document to the specified folder and file
        doc.save(output_file)
        print(f"Word document saved in chapter'{output_file}'")
        #print(p_tag)
    
    #find the url for next chap
        div_element = soup.find('div', class_='dis-hide', id='next_chp_url')

    # Extract the URL from the div element
        if div_element:
            url = div_element.text.strip()
            urls.pop()
            urls.append(url)
        else:
            print('URL not found')
    else:
        print("Failed to retrieve the webpage. Status code:", response.status_code)
